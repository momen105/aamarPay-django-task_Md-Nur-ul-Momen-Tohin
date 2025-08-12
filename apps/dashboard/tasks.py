from celery import shared_task
from .models import FileUpload, ActivityLog
from docx import Document
import os
import traceback

def get_text_from_docx(doc):
    texts = []
    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text)
    return "\n".join(texts)

@shared_task
def process_file_word_count(file_upload_id):
    try:
        file_upload = FileUpload.objects.get(id=file_upload_id)
        file_path = file_upload.file.path

        ext = os.path.splitext(file_path)[1].lower()
        word_count = 0

        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                word_count = len(text.split())

        elif ext == '.docx':
            doc = Document(file_path)
            full_text = get_text_from_docx(doc)
            word_count = len(full_text.split())

        else:
            file_upload.status = 'failed'
            file_upload.save()

            # Log failed activity with metadata
            ActivityLog.objects.create(
                user=file_upload.user,
                action="file_processing_failed",
                metadata={
                    "file_id": file_upload.id,
                    "filename": file_upload.filename,
                    "file_extension": ext,
                    "reason": "Unsupported file type"
                }
            )
            return

        # Update DB
        file_upload.word_count = word_count
        file_upload.status = 'completed'
        file_upload.save()

        # Activity log with metadata
        ActivityLog.objects.create(
            user=file_upload.user,
            action="file_processed",
            metadata={
                "file_id": file_upload.id,
                "filename": file_upload.filename,
                "file_extension": ext,
                "word_count": word_count,
                "file_size_bytes": file_upload.file.size,
                "storage_path": file_path
            }
        )

    except FileUpload.DoesNotExist:
        # Optional: log missing file case
        ActivityLog.objects.create(
            user=None,  # no user linked
            action="file_processing_failed",
            metadata={
                "file_id": file_upload_id,
                "reason": "FileUpload record not found"
            }
        )

    except Exception as e:
        # Mark as failed
        try:
            file_upload.status = 'failed'
            file_upload.save()
        except:
            pass

        # Detailed error log
        ActivityLog.objects.create(
            user=getattr(file_upload, "user", None),
            action="file_processing_failed",
            metadata={
                "file_id": file_upload_id,
                "error": str(e),
                "traceback": traceback.format_exc()
            }
        )
